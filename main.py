from docx import Document
import psycopg2
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
import datetime
import time
import os
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from contextlib import contextmanager
import schedule
import logging
import logging.handlers
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 尝试导入字体管理器，如果失败则提供备选方案
try:
    import matplotlib.font_manager as fm
    HAS_FONT_MANAGER = True
except ImportError:
    HAS_FONT_MANAGER = False
    import warnings
    warnings.warn("matplotlib.font_manager not available, using default fonts")

# 本地导入
from config.config import config
from webdav3.client import Client

# 类型别名
DatabaseResult = Tuple[List[str], List[Tuple]]

# 全局logger
logger = None

@dataclass
class ReportPeriod:
    """报告周期数据类"""
    start_date: datetime.date
    end_date: datetime.date
    start_timestamp: int
    end_timestamp: int
    
    @classmethod
    def create_weekly_period(cls) -> 'ReportPeriod':
        """创建一周的报告周期"""
        today = datetime.date.today()
        end_date = today - datetime.timedelta(days=1)
        start_date = end_date - datetime.timedelta(days=6)
        
        start_timestamp = int(time.mktime(start_date.timetuple()))
        end_timestamp = int(time.mktime(end_date.timetuple())) - 1
        
        return cls(
            start_date=start_date,
            end_date=end_date,
            start_timestamp=start_timestamp,
            end_timestamp=end_timestamp
        )


class DatabaseManager:
    """数据库管理类"""
    
    def __init__(self, connection_string: str):
        self.connection_string = connection_string
    
    @contextmanager
    def get_connection(self):
        """获取数据库连接上下文管理器"""
        conn = None
        try:
            conn = psycopg2.connect(self.connection_string)
            yield conn
        except Exception as e:
            logger.error(f"数据库连接失败: {e}")
            raise
        finally:
            if conn:
                conn.close()
    
    def execute_query(self, sql: str, params: Optional[tuple] = None) -> Optional[DatabaseResult]:
        """执行SQL查询"""
        try:
            with self.get_connection() as conn:
                with conn.cursor() as cursor:
                    logger.debug(f"执行SQL: {sql}")
                    if params:
                        cursor.execute(sql, params)
                    else:
                        cursor.execute(sql)
                    columns = [desc[0] for desc in cursor.description]
                    rows = cursor.fetchall()
                    return columns, rows
        except Exception as e:
            logger.error(f"查询失败: {e}")
            logger.error(f"SQL: {sql}")
            return None


class ReportGenerator:
    """报告生成器"""
    
    def __init__(self, config: Dict, period: ReportPeriod):
        self.config = config
        self.period = period
        self.db_manager = DatabaseManager(config['database_url'])
        self.doc = None
    
    def generate_report(self) -> Optional[str]:
        """生成报告主入口"""
        try:
            logger.info("开始生成报告")
            
            # 初始化文档
            self._init_document()
            
            # 生成报告各部分
            self._generate_summary()
            self._generate_protection_overview()
            self._generate_access_statistics()
            self._generate_attack_statistics()
            self._generate_detail_data()
            self._generate_report_info()
            
            # 保存文档
            filename = self._save_document()
            
            # 上传到WebDAV
            if self.config.get("webdav_options"):
                self._upload_to_webdav(filename)
            
            logger.info("报告生成完成")
            return filename
            
        except Exception as e:
            logger.error(f"报告生成失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    def _init_document(self):
        """初始化Word文档"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 设置标题样式
        if 'ReportHeading1' not in styles:
            style = styles.add_style('ReportHeading1', 1)
            style.font.name = '黑体'
            style.font.size = Pt(16)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 139)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        
        if 'ReportHeading2' not in styles:
            style = styles.add_style('ReportHeading2', 1)
            style.font.name = '黑体'
            style.font.size = Pt(13)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 100)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        
        # 设置正文样式
        if 'ReportBodyText' not in styles:
            style = styles.add_style('ReportBodyText', 1)
            style.font.name = '宋体'
            style.font.size = Pt(10.5)
            style.paragraph_format.line_spacing = 1.2
            style.paragraph_format.first_line_indent = Pt(21)
        
        # 设置强调样式
        if 'EmphasisText' not in styles:
            style = styles.add_style('EmphasisText', 2)
            style.font.name = '楷体'
            style.font.italic = True
            style.font.color.rgb = RGBColor(178, 34, 34)
    
    def _generate_summary(self):
        """生成摘要部分"""
        self.doc.add_heading("一、巡检必要性概述", level=1)
        text = (
            "本周对Web应用防火墙（WAF）进行了系统性巡检，旨在及时发现并处置潜在的安全威胁，"
            "确保Web应用服务的持续可用性和数据安全性。通过定期巡检，可有效识别异常攻击流量、"
            "优化防护策略、验证防护效果，并为安全态势评估提供数据支撑，降低因Web应用漏洞导致"
            "的数据泄露和服务中断风险。"
        )
        self.doc.add_paragraph(text, style='ReportBodyText')
    
    def _generate_protection_overview(self):
        """生成防护概览部分"""
        self.doc.add_heading("二、防护应用概览", level=1)
        
        # 获取总体数据
        summary_data = self._get_summary_data()
        if summary_data:
            template = (
                f"本周WAF总体运行平稳，总访问次数为:em{summary_data['访问总数']}:em，"
                f"总拦截次数为:em{summary_data['拦截总数']}:em，黑名单拦截次数为:em{summary_data['黑名单拦截数']}:em，"
                f"未拦截攻击次数为:em{summary_data['未拦截数']}:em，"
                f"拦截率为:em{self._calculate_intercept_rate(summary_data)}:em%。"
            )
            self._add_formatted_paragraph(template)
        else:
            self.doc.add_paragraph("无法获取总体统计数据", style='ReportBodyText')
        
        # 分应用查看
        self.doc.add_heading("2.1 分应用查看", level=2)
        self._get_protected_applications()
    
    def _add_formatted_paragraph(self, template: str) -> Optional[Any]:
        """添加格式化段落"""
        try:
            paragraph = self.doc.add_paragraph("", style='ReportBodyText')
            if not template:
                return paragraph
            
            parts = template.split(":em")
            
            for i, part in enumerate(parts):
                if part:  # 只添加非空文本
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # 奇数部分是强调文本
                        run.style = 'EmphasisText'
            
            return paragraph
        except Exception as e:
            logger.error(f"添加格式化段落失败: {e}")
            return None
    
    def _render_table(self, columns: List[str], rows: List[Tuple]):
        """渲染表格"""
        if not rows:
            self.doc.add_paragraph("暂无数据。", style='ReportBodyText')
            return
        
        table = self.doc.add_table(1, len(columns))
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, column in enumerate(columns):
            header_cells[i].text = str(column)
        
        # 添加数据行
        for row in rows:
            cells = table.add_row().cells
            for j, value in enumerate(row):
                cells[j].text = str(value)
    
    def _get_attack_type_name(self, attack_type: int) -> str:
        """获取攻击类型名称"""
        return self.config.get("attack_type_dict", {}).get(f'attack.type.{attack_type}', "未知攻击类型")
    
    def _convert_attack_types_in_rows(self, rows: List[Tuple], type_index: int) -> List[List]:
        """转换行数据中的攻击类型"""
        result = []
        for row in rows:
            row_list = list(row)
            row_list[type_index] = self._get_attack_type_name(row_list[type_index])
            result.append(row_list)
        return result
    
    def _build_except_clause(self, field: str, values: List[str]) -> str:
        if not values:
            return ""
        
        quoted_values = [f"{v}" for v in values]
        return f"and {field} not in ({','.join(quoted_values)})"
    
    def _get_summary_data(self) -> Optional[Dict]:
        """获取总体统计数据"""
        sql = f"""
        SELECT
            COALESCE(SUM(CASE WHEN mss."type" = 'website-req' THEN mss.value END)::int, 0) as 访问总数,
            COALESCE(SUM(CASE WHEN mss."type" = 'website-denied' THEN mss.value END)::int, 0) as 拦截总数,
            (
                SELECT COUNT(*) as 黑名单拦截数
                FROM mgt_rule_detect_log_basic mrdlb
                WHERE mrdlb.attack_type = -3
                AND mrdlb."timestamp" BETWEEN %s AND %s
                {self._build_except_clause('mrdlb.site_uuid', self.config.get('except_app_ids', []))}
            ),
            (
                SELECT COUNT(*) as 未拦截数
                FROM mgt_detect_log_basic mdlb
                WHERE mdlb."action" = 0
                AND mdlb."timestamp" BETWEEN %s AND %s
                {self._build_except_clause('mdlb.site_uuid', self.config.get('except_app_ids', []))}
            )
        FROM mgt_system_statistics mss
        WHERE mss.created_at BETWEEN %s AND %s
        {self._build_except_clause('mss.website', self.config.get('except_app_ids', []))}
        """
        
        params = (
            self.period.start_timestamp, self.period.end_timestamp,
            self.period.start_timestamp, self.period.end_timestamp,
            self.period.start_date, self.period.end_date
        )
        
        result = self.db_manager.execute_query(sql, params)
        if result:
            columns, rows = result
            if rows:
                return dict(zip(columns, rows[0]))
        return None
    
    def _calculate_intercept_rate(self, data: Dict) -> str:
        """计算拦截率"""
        if data['未拦截数'] == 0:
            return "100.00"
        total = data['拦截总数'] + data['未拦截数']
        rate = (data['拦截总数'] / total) * 100
        return f"{rate:.2f}"
    
    def _get_protected_applications(self):
        """获取防护应用列表"""
        sql = f"""
        SELECT
            mw.id as 应用序号,
            mw."comment" as 应用名称,
            mw.server_names as 域名,
            mw.ports as 开放端口,
            COALESCE(SUM(CASE WHEN mss."type" = 'website-req' THEN mss.value END)::int, 0) as 请求次数,
            COALESCE(SUM(CASE WHEN mss."type" = 'website-denied' THEN mss.value END)::int, 0) as 拦截次数
        FROM mgt_website mw
        LEFT JOIN mgt_system_statistics mss ON mw.id = mss.website::bigint
            AND mss.created_at BETWEEN %s AND %s
        {self._build_except_clause('mw.id', self.config.get('except_app_ids', []))}
        GROUP BY mw.id, mw."comment", mw.server_names, mw.ports
        ORDER BY mw.id
        """
        
        params = (self.period.start_date, self.period.end_date)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("暂无防护应用。", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询防护应用数据失败。", style='ReportBodyText')
    
    def _get_access_by_geography(self):
        """按地理区域获取访问数据"""
        sql = """
        SELECT
            country as 国家代号,
            province as 省份,
            city as 城市,
            SUM(count) as 访问次数
        FROM statistics_geos sg
        WHERE "time" BETWEEN %s AND %s
        GROUP BY country, province, city
        ORDER BY 访问次数 DESC
        LIMIT 10
        """
        
        params = (self.period.start_timestamp, self.period.end_timestamp)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                template = (
                    f"本周访问数据主要来自:em{rows[0][1]}:em:em{rows[0][2]}:em，"
                    f"访问次数为:em{rows[0][3]}:em，具体数据可参看下表。"
                )
                self._add_formatted_paragraph(template)
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("本周暂无访问数据。", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询地理访问数据失败。", style='ReportBodyText')
    
    def _get_access_by_ip(self):
        """按IP获取访问数据"""
        sql = f"""
        SELECT 
            si."key" as 访问IP,
            si.attack_type as 访问类型,
            SUM(si.count) as 访问次数
        FROM statistics_ips si
        WHERE si."time" BETWEEN %s AND %s
            AND si.attack_type = -1
            {self._build_except_clause('si.key', self.config.get('except_ips', []))}
        GROUP BY si."key", si.attack_type
        ORDER BY 访问次数 DESC
        LIMIT 10
        """
        
        params = (self.period.start_timestamp, self.period.end_timestamp)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                rows = self._convert_attack_types_in_rows(rows, 1)
                template = (
                    f"本周主要访问IP为:em{rows[0][0]}:em，"
                    f"访问次数为:em{rows[0][2]}:em，具体数据可参看下表。"
                )
                self._add_formatted_paragraph(template)
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("本周暂无访问数据。", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询IP访问数据失败。", style='ReportBodyText')
    
    def _get_attacks_by_type(self):
        """按类型获取攻击数据"""
        sql = f"""
        SELECT 
            si.attack_type as 攻击类型,
            SUM(si.count)::int as 攻击次数
        FROM statistics_ips si
        WHERE si."time" BETWEEN %s AND %s
            AND si.attack_type > 0
            {self._build_except_clause('si.key', self.config.get('except_ips', []))}
        GROUP BY si.attack_type
        ORDER BY 攻击次数 DESC
        """
        
        params = (self.period.start_timestamp, self.period.end_timestamp)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                rows = self._convert_attack_types_in_rows(rows, 0)
                
                # 生成饼图
                chart_path = self._create_attack_type_chart(rows)
                
                # 添加描述文本
                template = (
                    f"本周的主要攻击类型为:em{rows[0][0]}:em，"
                    f"该类型总计攻击:em{rows[0][1]}:em次，具体数据如下图表所示。"
                )
                paragraph = self._add_formatted_paragraph(template)
                
                # 添加图表（如果图表创建成功）
                if chart_path and os.path.exists(chart_path):
                    try:
                        # 添加一个空段落用于放置图片
                        picture_paragraph = self.doc.add_paragraph()
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = picture_paragraph.add_run()
                        run.add_picture(chart_path, width=Inches(6))
                        
                        # 清理临时文件
                        try:
                            os.remove(chart_path)
                        except:
                            pass
                    except Exception as e:
                        logger.error(f"添加图表失败: {e}")
                        self.doc.add_paragraph(f"图表生成失败: {str(e)}", style='ReportBodyText')
                
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("本周暂无攻击数据，您的WAF很安全", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询攻击类型数据失败。", style='ReportBodyText')
    
    def _create_attack_type_chart(self, rows: List[List]) -> Optional[str]:
        """创建攻击类型饼图"""
        try:
            if not rows or len(rows) == 0:
                logger.warning("没有数据可用于生成图表")
                return None
            
            # 提取数据和标签
            labels = [str(row[0])[:20] for row in rows]  # 限制标签长度
            values = [row[1] for row in rows]
            
            # 设置中文字体
            self._setup_matplotlib_font()
            
            # 创建饼图
            plt.figure(figsize=(8, 6), dpi=100)
            
            # 计算突出效果
            explode = [0.01] * len(values)
            if len(explode) > 0:
                explode[0] = 0.1
            
            # 绘制饼图
            wedges, texts, autotexts = plt.pie(
                values, 
                labels=labels, 
                explode=explode, 
                autopct='%.2f%%',
                startangle=90,
                shadow=False,
                textprops={'fontsize': 10}
            )
            
            # 设置百分比文本颜色
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontsize(9)
            
            plt.title("攻击类型统计图", fontsize=14, fontweight='bold')
            plt.axis('equal')  # 确保饼图是圆形
            
            # 保存图表
            chart_dir = Path("./charts")
            chart_dir.mkdir(exist_ok=True)
            
            chart_path = chart_dir / "attack_type_chart.png"
            plt.tight_layout()
            plt.savefig(chart_path, bbox_inches='tight', dpi=150)
            plt.close()
            
            logger.info(f"图表已保存: {chart_path}")
            return str(chart_path)
            
        except Exception as e:
            logger.error(f"创建图表失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    def _setup_matplotlib_font(self):
        """设置matplotlib字体"""
        try:
            # 尝试查找中文字体
            chinese_fonts = [
                '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                '/usr/share/fonts/wenquanyi/wqy-microhei/wqy-microhei.ttc',
                '/usr/share/fonts/truetype/arphic/uming.ttc',
                '/usr/share/fonts/truetype/arphic/ukai.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
            ]
            
            # 检查是否有可用的中文字体
            available_fonts = []
            for font_path in chinese_fonts:
                if os.path.exists(font_path):
                    available_fonts.append(font_path)
            
            if available_fonts and HAS_FONT_MANAGER:
                # 使用第一个可用的中文字体
                font_path = available_fonts[0]
                fm.fontManager.addfont(font_path)
                font_name = fm.FontProperties(fname=font_path).get_name()
                matplotlib.rcParams['font.sans-serif'] = [font_name]
                matplotlib.rcParams['axes.unicode_minus'] = False
                logger.info(f"使用中文字体: {font_name}")
                return True
            elif HAS_FONT_MANAGER:
                # 尝试使用系统字体
                system_fonts = ['DejaVu Sans', 'Arial', 'Helvetica', 'sans-serif']
                for font in system_fonts:
                    if font in fm.fontManager.ttflist:
                        matplotlib.rcParams['font.sans-serif'] = [font]
                        matplotlib.rcParams['axes.unicode_minus'] = False
                        logger.info(f"使用系统字体: {font}")
                        return True
            
            # 最后的回退方案
            matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans']
            matplotlib.rcParams['axes.unicode_minus'] = False
            logger.warning("使用默认字体，中文字符可能显示为方块")
            
            # 禁用警告，因为我们已经知道中文字体可能不可用
            import warnings
            warnings.filterwarnings('ignore', category=UserWarning)
            
            return False
            
        except Exception as e:
            logger.error(f"设置字体失败: {e}")
            # 设置默认字体
            matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans']
            matplotlib.rcParams['axes.unicode_minus'] = False
            
            # 禁用警告
            import warnings
            warnings.filterwarnings('ignore', category=UserWarning)
            
            return False
    
    def _get_attacks_by_ip(self):
        """按IP获取攻击数据"""
        sql = f"""
        SELECT 
            si."key" as 攻击IP,
            si.attack_type as 攻击类型,
            SUM(si.count) as 攻击次数
        FROM statistics_ips si
        WHERE si."time" BETWEEN %s AND %s
            AND si.attack_type > 0
            {self._build_except_clause('si.key', self.config.get('except_ips', []))}
        GROUP BY si."key", si.attack_type
        ORDER BY 攻击次数 DESC
        LIMIT 10
        """
        
        params = (self.period.start_timestamp, self.period.end_timestamp)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                rows = self._convert_attack_types_in_rows(rows, 1)
                template = (
                    f"本周的攻击主要来自:em{rows[0][0]}:em，"
                    f"攻击类型为:em{rows[0][1]}:em，"
                    f"总计攻击:em{rows[0][2]}:em次，具体数据参看下表。"
                )
                self._add_formatted_paragraph(template)
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("本周暂无攻击数据，您的WAF很安全", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询攻击IP数据失败。", style='ReportBodyText')
    
    def _get_unintercepted_attacks(self):
        """获取未拦截攻击明细"""
        sql = f"""
        SELECT 
            mw."comment" as 被攻击应用,
            mdlb.src_ip as 源IP,
            mdlb.host as 目标主机,
            mdlb.url_path as 请求路径,
            mdlb.dst_port as 目标端口,
            mdlb.country as 国家代码,
            mdlb.province as 省份,
            mdlb.city as 城市,
            mdlb.attack_type as 攻击类型,
            mdlb.updated_at as 攻击时间
        FROM mgt_detect_log_basic mdlb
        JOIN mgt_website mw ON mdlb.site_uuid::int = mw.id::int
        WHERE mdlb."timestamp" BETWEEN %s AND %s
            AND mdlb."action" = 0
            {self._build_except_clause('mdlb.site_uuid', self.config.get('except_app_ids', []))}
        """
        
        params = (self.period.start_timestamp, self.period.end_timestamp)
        result = self.db_manager.execute_query(sql, params)
        
        if result:
            columns, rows = result
            if rows:
                rows = self._convert_attack_types_in_rows(rows, 8)
                template = f"本周攻击有:em{len(rows)}:em条攻击未被拦截，我们将对其进行分析和拦截处理，具体数据参看下表。"
                self._add_formatted_paragraph(template)
                self._render_table(columns, rows)
            else:
                self.doc.add_paragraph("本周暂无未拦截攻击，所有攻击都被拒之门外。", style='ReportBodyText')
        else:
            self.doc.add_paragraph("查询未拦截攻击数据失败。", style='ReportBodyText')
    
    def _generate_access_statistics(self):
        """生成访问统计部分"""
        self.doc.add_heading("三、访问数据统计", level=1)
        
        # 按地理区域统计
        self.doc.add_heading("3.1 按地理区域统计访问数据TOP10", level=2)
        self._get_access_by_geography()
        
        # 按IP统计
        self.doc.add_heading("3.2 按访问IP统计访问数据TOP10", level=2)
        self._get_access_by_ip()
    
    def _generate_attack_statistics(self):
        """生成攻击统计部分"""
        self.doc.add_heading("四、攻击数据统计", level=1)
        
        # 按攻击方式统计
        self.doc.add_heading("4.1 按攻击方式统计数据", level=2)
        self._get_attacks_by_type()
        
        # 按攻击IP统计
        self.doc.add_heading("4.2 按攻击IP统计访问数据TOP10", level=2)
        self._get_attacks_by_ip()
    
    def _generate_detail_data(self):
        """生成明细数据部分"""
        self.doc.add_heading("五、明细数据展示", level=1)
        self.doc.add_heading("5.1 未拦截攻击明细", level=2)
        self._get_unintercepted_attacks()
    
    def _generate_report_info(self):
        """生成报告信息部分"""
        self.doc.add_heading("六、报告信息", level=1)
        
        info_items = [
            f"项目名称：{self.config['project_name']}",
            f"报告数据统计周期：{self.period.start_date}至{self.period.end_date}",
            f"报告生成时间：{datetime.datetime.now()}",
            f"报告审核人：{self.config['report_onwer']}"
        ]
        
        for item in info_items:
            self.doc.add_paragraph(item, style='ReportBodyText')
    
    def _save_document(self) -> str:
        """保存文档到本地"""
        # 确保报告目录存在
        report_dir = Path("./report")
        report_dir.mkdir(exist_ok=True)
        
        # 生成文件名
        filename = f"{self.config['project_name']}_{self.period.start_date}至{self.period.end_date}安全运维周报.docx"
        filepath = report_dir / filename
        
        # 保存文档
        self.doc.save(str(filepath))
        logger.info(f"文档已保存: {filepath}")
        
        return str(filepath)
    
    def _upload_to_webdav(self, local_path: str):
        """上传文档到WebDAV"""
        try:
            if not os.path.exists(local_path):
                logger.error(f"本地文件不存在: {local_path}")
                return
            
            client = Client(self.config["webdav_options"])
            
            # 创建远程目录
            remote_base = f"report/{self.config['report_onwer']}/{datetime.date.today().strftime('%Y%m%d')}"
            try:
                client.mkdir(remote_base)
            except:
                pass  # 目录可能已存在
            
            # 上传文件
            remote_path = f"{remote_base}/{os.path.basename(local_path)}"
            client.upload_sync(remote_path=remote_path, local_path=local_path)
            
            logger.info(f"文件已上传到WebDAV: {remote_path}")
            
        except Exception as e:
            logger.error(f"WebDAV上传失败: {e}")


class TaskScheduler:
    """任务调度器"""
    
    def __init__(self, config: Dict):
        self.config = config
    
    def run_immediately(self):
        """立即运行任务"""
        logger.info("立即执行生成报告")
        
        period = ReportPeriod.create_weekly_period()
        generator = ReportGenerator(self.config, period)
        generator.generate_report()
    
    def schedule_daily_task(self, time_str: str = "12:00"):
        """安排每日任务"""
        schedule.every().day.at(time_str).do(self._scheduled_task)
        
        logger.info(f"已安排每日任务，执行时间: {time_str}")
        
        # 主循环
        while True:
            schedule.run_pending()
            time.sleep(60)  # 每分钟检查一次
    
    def _scheduled_task(self):
        """定时任务"""
        logger.info("定时任务触发")
        self.run_immediately()


def setup_logging() -> logging.Logger:
    """设置日志配置"""
    # 创建日志目录
    log_dir = Path("./logs")
    log_dir.mkdir(exist_ok=True)
    
    # 创建logger
    logger_instance = logging.getLogger(__name__)
    logger_instance.setLevel(config.get('log_level', logging.INFO))
    
    # 如果已经配置过，直接返回
    if logger_instance.handlers:
        return logger_instance
    
    # 设置日志格式
    formatter = logging.Formatter(
        '%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    # 文件处理器（按大小轮转）
    file_handler = logging.handlers.RotatingFileHandler(
        './logs/app.log',
        maxBytes=100 * 1024 * 1024,  # 100MB
        backupCount=5,
        encoding='utf-8'
    )
    file_handler.setFormatter(formatter)
    logger_instance.addHandler(file_handler)
    
    # 控制台处理器
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    logger_instance.addHandler(console_handler)
    
    return logger_instance


def main():
    """主函数"""
    # 初始化日志
    global logger
    logger = setup_logging()
    
    # 创建调度器
    scheduler = TaskScheduler(config)
    
    # 根据命令行参数决定运行方式
    if len(sys.argv) == 2 and sys.argv[1] == '-now':
        scheduler.run_immediately()
    else:
        scheduler.schedule_daily_task()


if __name__ == '__main__':
    main()